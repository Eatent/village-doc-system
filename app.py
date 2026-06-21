# -*- coding: utf-8 -*-
"""
村委会文档智能分类与标注系统
作者：AI辅助生成
说明：本文件是整个系统的全部代码，只需要这一个 app.py + requirements.txt 即可运行
"""

import streamlit as st
import sqlite3
import re
import json
import io
from datetime import datetime

import pandas as pd

# ====== 可选依赖（用于解析不同文件类型）======
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

import requests

# ========================================================================
# 一、基础配置
# ========================================================================

st.set_page_config(page_title="村委会文档智能分类系统", page_icon="📄", layout="wide")

DB_PATH = "data.db"

CATEGORIES = ["政务通知类", "村民诉求与信访", "民政与福利", "安全与环保", "财务与资产", "其他/未分类"]
URGENCY_LEVELS = ["高", "中", "低"]
EVENT_TYPES = ["通知", "投诉/反映", "申请", "汇报/检查", "会议", "合同/账目", "其他"]
STATUS_LIST = ["未处理", "处理中", "已完成"]


# ========================================================================
# 二、数据库初始化与读写
# ========================================================================

def get_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    return conn


def init_db():
    conn = get_conn()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            upload_time TEXT,
            file_name TEXT,
            raw_content TEXT,
            summary TEXT,
            category TEXT,
            persons TEXT,
            key_date TEXT,
            event_type TEXT,
            urgency TEXT,
            status TEXT
        )
    """)
    conn.commit()
    conn.close()


def save_record(record: dict):
    conn = get_conn()
    conn.execute("""
        INSERT INTO documents
        (upload_time, file_name, raw_content, summary, category, persons, key_date, event_type, urgency, status)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        record["upload_time"], record["file_name"], record["raw_content"], record["summary"],
        record["category"], record["persons"], record["key_date"], record["event_type"],
        record["urgency"], record["status"]
    ))
    conn.commit()
    conn.close()


def load_all_records() -> pd.DataFrame:
    conn = get_conn()
    df = pd.read_sql_query("SELECT * FROM documents ORDER BY id DESC", conn)
    conn.close()
    return df


def update_status(record_id: int, new_status: str):
    conn = get_conn()
    conn.execute("UPDATE documents SET status=? WHERE id=?", (new_status, record_id))
    conn.commit()
    conn.close()


def delete_record(record_id: int):
    conn = get_conn()
    conn.execute("DELETE FROM documents WHERE id=?", (record_id,))
    conn.commit()
    conn.close()


# ========================================================================
# 三、文档解析：把 Word / PDF / Excel 转成纯文本
# ========================================================================

def parse_docx(file) -> str:
    if DocxDocument is None:
        return ""
    doc = DocxDocument(file)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def parse_pdf(file) -> str:
    if PdfReader is None:
        return ""
    reader = PdfReader(file)
    text_parts = []
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def parse_xlsx(file) -> str:
    sheets = pd.read_excel(file, sheet_name=None, header=None)
    text_parts = []
    for sheet_name, df in sheets.items():
        text_parts.append(f"【表格：{sheet_name}】")
        for _, row in df.iterrows():
            row_text = " ".join(str(v) for v in row.values if pd.notna(v))
            if row_text.strip():
                text_parts.append(row_text)
    return "\n".join(text_parts)


def parse_uploaded_file(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".docx"):
        return parse_docx(uploaded_file)
    elif name.endswith(".pdf"):
        return parse_pdf(uploaded_file)
    elif name.endswith(".xlsx") or name.endswith(".xls"):
        return parse_xlsx(uploaded_file)
    else:
        return uploaded_file.read().decode("utf-8", errors="ignore")


# ========================================================================
# 四、分类逻辑（两种模式：① 调用AI接口  ② 本地规则兜底）
# ========================================================================

# ---- 4.1 本地规则关键词库（无需联网、不花钱，AI接口不可用时自动启用）----

CATEGORY_KEYWORDS = {
    "政务通知类": ["通知", "文件", "会议精神", "上级", "政策", "指标", "传达", "部署", "县政府", "镇政府", "街道办"],
    "村民诉求与信访": ["反映", "投诉", "信访", "漏水", "修路", "纠纷", "矛盾", "建议", "求助", "诉求", "上访"],
    "民政与福利": ["低保", "残疾人", "补助", "高龄津贴", "退役军人", "优抚", "救助", "民政", "五保户", "孤寡"],
    "安全与环保": ["消防", "防汛", "抗旱", "垃圾分类", "环境整治", "安全检查", "隐患", "环保", "巡查"],
    "财务与资产": ["集体经济", "报销", "账目", "土地流转", "合同", "收入", "支出", "资产", "审计", "经费"],
}

URGENT_KEYWORDS_HIGH = ["紧急", "立即", "限期", "24小时", "今日前", "马上", "刻不容缓", "重大隐患"]
URGENT_KEYWORDS_MID = ["尽快", "近期", "本周", "月底前", "请及时"]

EVENT_KEYWORDS = {
    "通知": ["通知", "传达", "部署"],
    "投诉/反映": ["投诉", "反映", "纠纷", "上访"],
    "申请": ["申请", "审批", "审核"],
    "汇报/检查": ["汇报", "检查", "巡查", "整治"],
    "会议": ["会议", "会议纪要", "研究决定"],
    "合同/账目": ["合同", "账目", "报销", "收入", "支出"],
}

DATE_PATTERN = re.compile(
    r"(\d{4}年\d{1,2}月\d{1,2}日|\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2}|"
    r"\d{1,2}月底前?|本周内|月底前|年底前)"
)

# 简单人名提取：抓“姓名/村民/反映人/申请人/联系人”等关键词后面的2-4个中文字
NAME_PATTERN = re.compile(
    r"(?:村民|反映人|申请人|联系人|姓名|当事人)[：:]?\s*([\u4e00-\u9fa5]{2,4})"
)


def rule_based_classify(text: str) -> dict:
    """本地规则分类：不联网、零成本、速度快，但精度不如AI接口"""
    scores = {cat: 0 for cat in CATEGORY_KEYWORDS}
    for cat, kws in CATEGORY_KEYWORDS.items():
        for kw in kws:
            scores[cat] += text.count(kw)
    best_category = max(scores, key=scores.get)
    if scores[best_category] == 0:
        best_category = "其他/未分类"

    # 紧急程度
    urgency = "低"
    if any(kw in text for kw in URGENT_KEYWORDS_HIGH):
        urgency = "高"
    elif any(kw in text for kw in URGENT_KEYWORDS_MID):
        urgency = "中"

    # 事件类型
    event_type = "其他"
    event_scores = {et: 0 for et in EVENT_KEYWORDS}
    for et, kws in EVENT_KEYWORDS.items():
        for kw in kws:
            event_scores[et] += text.count(kw)
    if max(event_scores.values()) > 0:
        event_type = max(event_scores, key=event_scores.get)

    # 日期与人名
    dates = DATE_PATTERN.findall(text)
    names = NAME_PATTERN.findall(text)

    summary = text.strip().replace("\n", " ")[:60] + ("..." if len(text) > 60 else "")

    return {
        "category": best_category,
        "urgency": urgency,
        "event_type": event_type,
        "persons": "、".join(sorted(set(names))) if names else "",
        "key_date": "、".join(sorted(set(dates))) if dates else "",
        "summary": summary,
    }


# ---- 4.2 AI接口分类（精度更高，需要在 Streamlit secrets 中配置密钥）----

def ai_classify(text: str) -> dict:
    """
    调用 DeepSeek 或 OpenAI 接口进行分类。
    需要在 .streamlit/secrets.toml 中配置：
        AI_PROVIDER = "deepseek"   # 或 "openai"
        AI_API_KEY = "你的密钥"
    若未配置，本函数不会被调用，系统自动使用本地规则分类。
    """
    provider = st.secrets.get("AI_PROVIDER", "")
    api_key = st.secrets.get("AI_API_KEY", "")
    if not provider or not api_key:
        return None

    if provider == "deepseek":
        url = "https://api.deepseek.com/chat/completions"
        model = "deepseek-chat"
    elif provider == "openai":
        url = "https://api.openai.com/v1/chat/completions"
        model = "gpt-4o-mini"
    else:
        return None

    prompt = f"""你是村委会的文档分类助手。请阅读以下文档内容，严格按JSON格式输出，不要任何多余文字、不要markdown代码块标记。

文档内容：
{text[:3000]}

请输出如下JSON结构：
{{
  "category": "从【政务通知类/村民诉求与信访/民政与福利/安全与环保/财务与资产/其他】中选一个",
  "urgency": "从【高/中/低】中选一个",
  "event_type": "从【通知/投诉/反映/申请/汇报/检查/会议/合同/账目/其他】中选一个",
  "persons": "文中出现的人名，用、分隔，没有则留空",
  "key_date": "文中出现的关键日期或期限，没有则留空",
  "summary": "一句话摘要，30字以内"
}}"""

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
    }

    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=30)
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        content = content.strip().strip("```json").strip("```").strip()
        result = json.loads(content)
        return result
    except Exception as e:
        st.warning(f"AI接口调用失败，已自动切换为本地规则分类。错误信息：{e}")
        return None


def classify_text(text: str) -> dict:
    ai_result = ai_classify(text)
    if ai_result:
        return ai_result
    return rule_based_classify(text)


# ========================================================================
# 五、导出 Excel
# ========================================================================

def export_to_excel(df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    export_df = df.rename(columns={
        "id": "编号", "upload_time": "录入时间", "file_name": "文件名",
        "summary": "内容摘要", "category": "业务板块", "persons": "涉及人名",
        "key_date": "涉及时间", "event_type": "事件类型", "urgency": "紧急程度", "status": "处理状态"
    })
    export_df = export_df.drop(columns=["raw_content"], errors="ignore")
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        export_df.to_excel(writer, index=False, sheet_name="文档台账")
    return output.getvalue()


# ========================================================================
# 六、页面：上传与分析
# ========================================================================

def page_upload():
    st.header("📤 上传文档 / 粘贴文本，自动分类")

    tab1, tab2 = st.tabs(["上传文件", "粘贴文本"])
    text_content = ""
    file_name = ""

    with tab1:
        uploaded_file = st.file_uploader("支持 Word(.docx) / PDF / Excel(.xlsx)", type=["docx", "pdf", "xlsx", "xls"])
        if uploaded_file:
            file_name = uploaded_file.name
            with st.spinner("正在解析文档..."):
                text_content = parse_uploaded_file(uploaded_file)
            if text_content.strip():
                st.success(f"已成功解析「{file_name}」，共 {len(text_content)} 字")
                with st.expander("查看解析出的原文内容"):
                    st.text(text_content[:2000])
            else:
                st.error("未能解析出文字内容，请检查文件格式或文件是否为扫描图片版PDF")

    with tab2:
        pasted = st.text_area("直接粘贴长文本", height=200, placeholder="将通知、信访记录、申请材料等文字粘贴到这里...")
        if pasted.strip():
            text_content = pasted
            file_name = "（粘贴文本）"

    if text_content.strip() and st.button("🚀 开始自动分析", type="primary"):
        with st.spinner("AI正在分析文档内容..."):
            result = classify_text(text_content)
        st.session_state["pending_text"] = text_content
        st.session_state["pending_filename"] = file_name
        st.session_state["pending_result"] = result

    # 展示分析结果，允许人工修正后保存
    if "pending_result" in st.session_state:
        st.divider()
        st.subheader("✅ AI识别结果（可手动修改后保存）")
        result = st.session_state["pending_result"]

        col1, col2 = st.columns(2)
        with col1:
            category = st.selectbox("业务板块", CATEGORIES,
                                     index=CATEGORIES.index(result.get("category", "其他/未分类"))
                                     if result.get("category") in CATEGORIES else len(CATEGORIES) - 1)
            event_type = st.selectbox("事件类型", EVENT_TYPES,
                                       index=EVENT_TYPES.index(result.get("event_type", "其他"))
                                       if result.get("event_type") in EVENT_TYPES else len(EVENT_TYPES) - 1)
            urgency = st.selectbox("紧急程度", URGENCY_LEVELS,
                                    index=URGENCY_LEVELS.index(result.get("urgency", "低"))
                                    if result.get("urgency") in URGENCY_LEVELS else 2)
        with col2:
            persons = st.text_input("涉及人名（用、分隔）", value=result.get("persons", ""))
            key_date = st.text_input("涉及时间", value=result.get("key_date", ""))
            summary = st.text_input("内容摘要", value=result.get("summary", ""))

        if st.button("💾 确认保存到台账"):
            record = {
                "upload_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "file_name": st.session_state["pending_filename"],
                "raw_content": st.session_state["pending_text"],
                "summary": summary,
                "category": category,
                "persons": persons,
                "key_date": key_date,
                "event_type": event_type,
                "urgency": urgency,
                "status": "未处理",
            }
            save_record(record)
            st.success("已保存！可前往左侧「台账列表」查看")
            del st.session_state["pending_result"]
            del st.session_state["pending_text"]
            del st.session_state["pending_filename"]


# ========================================================================
# 七、页面：台账列表 / 筛选 / 导出
# ========================================================================

def page_list():
    st.header("📋 文档台账列表")

    df = load_all_records()
    if df.empty:
        st.info("暂无数据，请先到「上传与分析」页面录入文档")
        return

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        f_category = st.selectbox("业务板块筛选", ["全部"] + CATEGORIES)
    with col2:
        f_urgency = st.selectbox("紧急程度筛选", ["全部"] + URGENCY_LEVELS)
    with col3:
        f_status = st.selectbox("处理状态筛选", ["全部"] + STATUS_LIST)
    with col4:
        keyword = st.text_input("关键词搜索（摘要/人名/原文）")

    filtered = df.copy()
    if f_category != "全部":
        filtered = filtered[filtered["category"] == f_category]
    if f_urgency != "全部":
        filtered = filtered[filtered["urgency"] == f_urgency]
    if f_status != "全部":
        filtered = filtered[filtered["status"] == f_status]
    if keyword.strip():
        mask = (filtered["summary"].str.contains(keyword, na=False) |
                filtered["persons"].str.contains(keyword, na=False) |
                filtered["raw_content"].str.contains(keyword, na=False))
        filtered = filtered[mask]

    st.caption(f"共 {len(filtered)} 条记录")

    display_cols = ["id", "upload_time", "file_name", "summary", "category",
                     "event_type", "urgency", "persons", "key_date", "status"]
    display_df = filtered[display_cols].rename(columns={
        "id": "编号", "upload_time": "录入时间", "file_name": "文件名", "summary": "摘要",
        "category": "板块", "event_type": "事件类型", "urgency": "紧急", "persons": "人名",
        "key_date": "时间", "status": "状态"
    })
    st.dataframe(display_df, use_container_width=True, height=400)

    st.divider()
    st.subheader("更新处理状态 / 删除记录")
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        record_id = st.number_input("输入编号", min_value=0, step=1)
    with col_b:
        new_status = st.selectbox("更新为", STATUS_LIST)
        if st.button("更新状态"):
            update_status(int(record_id), new_status)
            st.success("已更新，请刷新查看")
    with col_c:
        if st.button("🗑️ 删除该条记录", type="secondary"):
            delete_record(int(record_id))
            st.success("已删除，请刷新查看")

    st.divider()
    excel_data = export_to_excel(filtered)
    st.download_button(
        "📥 导出当前筛选结果为 Excel",
        data=excel_data,
        file_name=f"村委会文档台账_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ========================================================================
# 八、主程序入口
# ========================================================================

def main():
    init_db()
    st.sidebar.title("📄 村委会文档智能分类系统")
    page = st.sidebar.radio("功能导航", ["📤 上传与分析", "📋 台账列表与导出"])

    ai_status = "✅ 已配置AI接口" if st.secrets.get("AI_API_KEY", "") else "⚙️ 当前使用本地规则分类（未配置AI接口）"
    st.sidebar.caption(ai_status)
    st.sidebar.caption("数据已保存在本地数据库 data.db")

    if page == "📤 上传与分析":
        page_upload()
    else:
        page_list()


if __name__ == "__main__":
    main()
